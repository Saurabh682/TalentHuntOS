"""Local, bounded text extraction from recruiter-provided resume artifacts."""

from __future__ import annotations

import io
import re
from pathlib import Path


MAX_RESUME_BYTES = 8 * 1024 * 1024
MAX_RESUME_CHARS = 200_000
MAX_PDF_PAGES = 75
SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".docx", ".txt"}


class ResumeImportError(ValueError):
    pass


def _clean_text(value: str) -> str:
    value = (value or "").replace("\x00", "")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    return cleaned[:MAX_RESUME_CHARS]


def _extract_pdf(content: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise ResumeImportError("The PDF could not be read.") from exc
    if reader.is_encrypted:
        raise ResumeImportError("Password-protected PDFs are not supported.")
    page_count = min(len(reader.pages), MAX_PDF_PAGES)
    parts: list[str] = []
    for page in reader.pages[:page_count]:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return _clean_text("\n".join(parts)), page_count


def _extract_docx(content: bytes) -> tuple[str, int]:
    from docx import Document

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ResumeImportError("The DOCX file could not be read.") from exc
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return _clean_text("\n".join(parts)), len(document.paragraphs)


def _extract_txt(content: bytes) -> tuple[str, int]:
    try:
        value = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        value = content.decode("latin-1")
    return _clean_text(value), max(1, value.count("\n") + 1)


def extract_resume_artifact(filename: str, content: bytes) -> dict[str, object]:
    """Extract text in memory; no uploaded artifact is persisted to disk."""
    safe_name = Path(filename or "resume").name
    extension = Path(safe_name).suffix.casefold()
    if extension not in SUPPORTED_RESUME_EXTENSIONS:
        raise ResumeImportError("Use a PDF, DOCX, or TXT resume.")
    if not content:
        raise ResumeImportError("The uploaded resume is empty.")
    if len(content) > MAX_RESUME_BYTES:
        raise ResumeImportError("Resume files must be 8 MB or smaller.")

    if extension == ".pdf":
        text, units = _extract_pdf(content)
    elif extension == ".docx":
        text, units = _extract_docx(content)
    else:
        text, units = _extract_txt(content)
    if len(text) < 40:
        raise ResumeImportError(
            "The resume did not contain enough readable text. Scanned PDFs need OCR before import."
        )
    return {
        "filename": safe_name,
        "extension": extension,
        "size_bytes": len(content),
        "text": text,
        "text_chars": len(text),
        "units_read": units,
        "persisted": False,
    }


import io

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from app.candidates.resume_import import ResumeImportError, extract_resume_artifact


RESUME_TEXT = (
    "Nina Shah\nSpine Animator\nFive years creating 2D character animation.\n"
    "Skills: Spine, Photoshop, After Effects.\nExperience: Animator at Studio One."
)


def test_extracts_txt_resume_without_persisting_artifact():
    result = extract_resume_artifact("nina.txt", RESUME_TEXT.encode("utf-8"))
    assert result["text"].startswith("Nina Shah")
    assert result["extension"] == ".txt"
    assert result["persisted"] is False


def test_extracts_docx_and_pdf_resume_text():
    doc_buffer = io.BytesIO()
    document = Document()
    document.add_heading("Nina Shah", 1)
    document.add_paragraph(RESUME_TEXT)
    document.save(doc_buffer)
    doc_result = extract_resume_artifact("resume.docx", doc_buffer.getvalue())
    assert "Spine Animator" in doc_result["text"]

    pdf_buffer = io.BytesIO()
    pdf = canvas.Canvas(pdf_buffer)
    y = 800
    for line in RESUME_TEXT.splitlines():
        pdf.drawString(72, y, line)
        y -= 18
    pdf.save()
    pdf_result = extract_resume_artifact("resume.pdf", pdf_buffer.getvalue())
    assert "Photoshop" in pdf_result["text"]


def test_resume_import_rejects_unsupported_empty_and_oversized_files():
    with pytest.raises(ResumeImportError, match="PDF, DOCX, or TXT"):
        extract_resume_artifact("resume.exe", b"not a resume")
    with pytest.raises(ResumeImportError, match="empty"):
        extract_resume_artifact("resume.txt", b"")
    with pytest.raises(ResumeImportError, match="8 MB"):
        extract_resume_artifact("resume.txt", b"x" * (8 * 1024 * 1024 + 1))


def test_candidate_detail_upload_routes_through_reviewed_profile_action():
    from pathlib import Path

    root = Path(__file__).resolve().parents[1]
    detail = (root / "app/ui/pages/candidate_detail.py").read_text(encoding="utf-8")
    review = (root / "app/ui/components/profile_review_dialog.py").read_text(encoding="utf-8")
    assert "extract_resume_artifact" in detail
    assert "await event.file.read()" in detail
    assert "run_extract_then_review" in detail
    assert '"candidates.profile.apply"' in review
    assert '"resume_text": draft.get("resume_text")' in review
